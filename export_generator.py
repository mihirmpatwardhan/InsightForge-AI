import re
from io import BytesIO
from docx import Document
from docx.shared import RGBColor
from reportlab.lib.pagesizes import A4
from reportlab.lib import colors
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, HRFlowable


def add_formatted_docx_paragraph(doc, text: str, style=None):
    p = doc.add_paragraph(style=style)
    parts = re.split(r"(\*\*.*?\*\*)", text)
    for part in parts:
        if part.startswith("**") and part.endswith("**"):
            run = p.add_run(part[2:-2])
            run.bold = True
        else:
            p.add_run(part)
    return p


def create_docx_bytes(title: str, content: str) -> BytesIO:
    buffer = BytesIO()
    doc = Document()

    heading = doc.add_heading(level=0)
    run = heading.add_run(title)
    run.font.color.rgb = RGBColor(79, 70, 229)

    for line in content.split("\n"):
        clean_line = line.strip()
        if not clean_line:
            continue

        if clean_line.startswith("# "):
            doc.add_heading(clean_line[2:].replace("**", ""), level=1)
        elif clean_line.startswith("## "):
            doc.add_heading(clean_line[3:].replace("**", ""), level=2)
        elif clean_line.startswith("### "):
            doc.add_heading(clean_line[4:].replace("**", ""), level=3)
        elif clean_line.startswith(("- ", "* ")):
            add_formatted_docx_paragraph(doc, clean_line[2:], style="List Bullet")
        elif re.match(r"^\d+\.\s", clean_line):
            sub_text = re.sub(r"^\d+\.\s", "", clean_line)
            add_formatted_docx_paragraph(doc, sub_text, style="List Number")
        else:
            add_formatted_docx_paragraph(doc, clean_line)

    doc.save(buffer)
    buffer.seek(0)
    return buffer


def md_to_html(text: str) -> str:
    text = re.sub(r"\*\*(.*?)\*\*", r"<b>\1</b>", text)
    text = re.sub(r"&(?!#?\w+;)", "&amp;", text)
    return text


def create_pdf_bytes(title: str, content: str) -> BytesIO:
    buffer = BytesIO()
    doc = SimpleDocTemplate(
        buffer,
        pagesize=A4,
        rightMargin=40,
        leftMargin=40,
        topMargin=40,
        bottomMargin=40
    )

    styles = getSampleStyleSheet()
    title_style = ParagraphStyle(
        "DocTitle",
        parent=styles["Heading1"],
        fontName="Helvetica-Bold",
        fontSize=20,
        leading=24,
        textColor=colors.HexColor("#4f46e5"),
        spaceAfter=12
    )
    h1_style = ParagraphStyle(
        "Header1",
        parent=styles["Heading2"],
        fontName="Helvetica-Bold",
        fontSize=13,
        leading=17,
        textColor=colors.HexColor("#1e1b4b"),
        spaceBefore=10,
        spaceAfter=4
    )
    body_style = ParagraphStyle(
        "BodyTextCustom",
        parent=styles["Normal"],
        fontName="Helvetica",
        fontSize=9.5,
        leading=13.5,
        textColor=colors.HexColor("#334155"),
        spaceAfter=5
    )
    bullet_style = ParagraphStyle(
        "BulletCustom",
        parent=body_style,
        leftIndent=14,
        spaceAfter=3
    )

    story = [
        Paragraph(md_to_html(title), title_style),
        HRFlowable(width="100%", thickness=1, color=colors.HexColor("#6366f1"), spaceAfter=12)
    ]

    for line in content.split("\n"):
        clean_line = line.strip()
        if not clean_line:
            story.append(Spacer(1, 3))
            continue

        html_line = md_to_html(clean_line)
        if clean_line.startswith(("# ", "## ", "### ")):
            header_text = re.sub(r"^#+\s*", "", html_line)
            story.append(Paragraph(header_text, h1_style))
        elif clean_line.startswith(("- ", "* ")):
            bullet_text = f"• {re.sub(r'^[-\*]\s*', '', html_line)}"
            story.append(Paragraph(bullet_text, bullet_style))
        elif re.match(r"^\d+\.\s", clean_line):
            story.append(Paragraph(html_line, bullet_style))
        else:
            story.append(Paragraph(html_line, body_style))

    doc.build(story)
    buffer.seek(0)
    return buffer